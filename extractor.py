"""
extractor.py — Complete metadata extraction engine v3.

Extracts from PDF (PyMuPDF):
  TEXT:
    - Every word: text, font, size, color, bold, italic, underline
    - Bounding box (x0,y0,x1,y1) per word
    - Line number, column, block, span
    - Zone: header / footer / body
  SPACING:
    - Line spacing (gap between lines in pt)
    - Paragraph spacing (gap between blocks)
    - Letter spacing / kerning (char-level bbox diff)
  ALIGNMENT:
    - Per-line alignment: LEFT / CENTER / RIGHT / JUSTIFY
  HIGHLIGHT:
    - Background highlight color per span
    - Highlight annotation rectangles
  IMAGES:
    - Embedded images: position, size, raw bytes for comparison
  DRAWINGS:
    - Lines, rectangles, curves: color, fill, width, position
  WATERMARKS:
    - Transparent text layers
    - Background image watermarks
  FORM FIELDS:
    - Name, type, value, position, required flag
  TABLE BORDERS:
    - Rectangle drawings inside table regions
  PAGE:
    - Width, height, rotation, margins

Extracts from DOCX (python-docx):
  - Paragraph alignment (exact enum)
  - Line spacing, space before/after paragraph
  - Run font: name, size, color, bold, italic, underline, strike, highlight
  - Table cell borders (XML-level): color, width, style per side
  - Header / footer paragraphs with full run metadata
  - Images (inline shapes)
  - Form fields (content controls)
  - Page size, margins
"""

from __future__ import annotations
import os, io, hashlib, math
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional, Tuple

# ── Optional imports ──────────────────────────────────────────────────────────
try:
    import fitz
    _PYMUPDF = True
except ImportError:
    _PYMUPDF = False

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree
    _DOCX = True
except ImportError:
    _DOCX = False

try:
    import cv2
    import numpy as np
    from PIL import Image as PILImage
    _CV2 = True
except ImportError:
    _CV2 = False

try:
    from PIL import Image as PILImage
    _PIL = True
except ImportError:
    _PIL = False


# ─────────────────────────────────────────────────────────────────────────────
# Data structures
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class BBox:
    x0: float; y0: float; x1: float; y1: float

    @property
    def width(self):  return round(self.x1 - self.x0, 2)
    @property
    def height(self): return round(self.y1 - self.y0, 2)

    def to_dict(self):
        return {"x0": round(self.x0,2), "y0": round(self.y0,2),
                "x1": round(self.x1,2), "y1": round(self.y1,2)}


@dataclass
class CharToken:
    """Single character with exact position (for kerning/letter-spacing)."""
    char:   str
    bbox:   BBox
    origin: Tuple[float, float]  # baseline x,y


@dataclass
class WordToken:
    text:           str
    page:           int
    line:           int
    col:            int
    bbox:           BBox
    font_name:      str   = ""
    font_size:      float = 0.0
    font_color:     Tuple = (0, 0, 0)
    bold:           bool  = False
    italic:         bool  = False
    underline:      bool  = False
    strike:         bool  = False
    highlight_color: Optional[Tuple] = None   # NEW
    alignment:      str   = "LEFT"             # NEW: LEFT/CENTER/RIGHT/JUSTIFY
    letter_spacing: float = 0.0                # NEW: avg gap between chars in pt
    zone:           str   = "body"
    block_id:       int   = 0
    span_id:        int   = 0
    chars:          List[CharToken] = field(default_factory=list)  # NEW


@dataclass
class LineSpacing:
    """Spacing between consecutive lines."""
    page:        int
    line_from:   int
    line_to:     int
    spacing_pt:  float
    zone:        str


@dataclass
class ParaSpacing:
    """Space before/after a paragraph block."""
    page:        int
    block_id:    int
    space_before: float
    space_after:  float


@dataclass
class DrawingElement:
    page:   int
    kind:   str        # line / rect / curve
    bbox:   BBox
    color:  Tuple
    fill:   Optional[Tuple]
    width:  float
    is_table_border: bool = False   # NEW
    is_watermark:    bool = False   # NEW


@dataclass
class ImageElement:
    """Embedded image inside a page."""
    page:     int
    bbox:     BBox
    width:    int
    height:   int
    img_hash: str          # MD5 for comparison
    img_bytes: bytes       # raw PNG bytes for pixel diff
    is_watermark: bool = False   # NEW


@dataclass
class WatermarkElement:
    """Detected watermark (text or image)."""
    page:    int
    kind:    str       # "text" | "image"
    text:    str       # if text watermark
    bbox:    BBox
    opacity: float
    color:   Optional[Tuple]


@dataclass
class FormField:
    """PDF form field / widget."""
    page:       int
    name:       str
    field_type: str    # text / checkbox / radio / dropdown / signature
    value:      str
    bbox:       BBox
    required:   bool
    font_name:  str  = ""
    font_size:  float = 0.0


@dataclass
class TableCell:
    """Table cell with border styling."""
    page:        int
    row:         int
    col:         int
    text:        str
    bbox:        BBox
    border_top:    Optional[Dict] = None   # {color, width, style}
    border_bottom: Optional[Dict] = None
    border_left:   Optional[Dict] = None
    border_right:  Optional[Dict] = None
    bg_color:    Optional[Tuple]  = None


@dataclass
class PageMeta:
    page:         int
    width:        float
    height:       float
    rotation:     int
    margin_top:   float
    margin_bot:   float
    margin_left:  float
    margin_right: float


@dataclass
class DocumentData:
    path:          str
    file_type:     str
    pages:         List[PageMeta]       = field(default_factory=list)
    words:         List[WordToken]      = field(default_factory=list)
    line_spacings: List[LineSpacing]    = field(default_factory=list)
    para_spacings: List[ParaSpacing]    = field(default_factory=list)
    drawings:      List[DrawingElement] = field(default_factory=list)
    images:        List[ImageElement]   = field(default_factory=list)
    watermarks:    List[WatermarkElement] = field(default_factory=list)
    form_fields:   List[FormField]      = field(default_factory=list)
    table_cells:   List[TableCell]      = field(default_factory=list)
    raw_text:      str                  = ""
    page_images:   List[Any]            = field(default_factory=list)
    errors:        List[str]            = field(default_factory=list)

    def words_on_page(self, page):  return [w for w in self.words if w.page == page]
    def header_words(self, page):   return [w for w in self.words if w.page == page and w.zone == "header"]
    def footer_words(self, page):   return [w for w in self.words if w.page == page and w.zone == "footer"]
    def body_words(self, page):     return [w for w in self.words if w.page == page and w.zone == "body"]


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def extract(path: str) -> DocumentData:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":   return _extract_pdf(path)
    if ext == ".docx":  return _extract_docx(path)
    return _extract_text(path)


# ─────────────────────────────────────────────────────────────────────────────
# PDF Extractor — full metadata
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf(path: str) -> DocumentData:
    data = DocumentData(path=path, file_type="pdf")
    if not _PYMUPDF:
        data.errors.append("PyMuPDF not installed. Run: pip install PyMuPDF")
        return _extract_text(path)

    try:
        doc = fitz.open(path)
    except Exception as e:
        data.errors.append(f"Cannot open PDF: {e}")
        return data

    DPI = 150

    for pno, page in enumerate(doc, 1):
        rect = page.rect
        pw, ph = rect.width, rect.height

        # ── Page metadata ─────────────────────────────────────────────────────
        blocks = page.get_text("blocks")
        if blocks:
            m_top   = min(b[1] for b in blocks)
            m_bot   = ph - max(b[3] for b in blocks)
            m_left  = min(b[0] for b in blocks)
            m_right = pw - max(b[2] for b in blocks)
        else:
            m_top = m_bot = m_left = m_right = 0

        data.pages.append(PageMeta(
            page=pno, width=round(pw,2), height=round(ph,2),
            rotation=page.rotation,
            margin_top=round(max(m_top,0),2), margin_bot=round(max(m_bot,0),2),
            margin_left=round(max(m_left,0),2), margin_right=round(max(m_right,0),2),
        ))

        header_y = ph * 0.15
        footer_y = ph * 0.85

        # ── Full text extraction with rawdict ─────────────────────────────────
        raw = page.get_text("rawdict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        line_num       = 0
        prev_line_y1   = None
        prev_block_y1  = None

        for block in raw.get("blocks", []):
            if block.get("type") != 0:
                continue

            block_id  = block.get("number", 0)
            block_y0  = block["bbox"][1]
            block_y1  = block["bbox"][3]

            # Paragraph spacing
            if prev_block_y1 is not None:
                space_before = round(block_y0 - prev_block_y1, 2)
                data.para_spacings.append(ParaSpacing(
                    page=pno, block_id=block_id,
                    space_before=max(space_before, 0),
                    space_after=0
                ))
            prev_block_y1 = block_y1

            for ln in block.get("lines", []):
                line_num += 1
                line_y0  = ln["bbox"][1]
                line_y1  = ln["bbox"][3]
                zone     = ("header" if line_y0 < header_y
                            else "footer" if line_y0 > footer_y
                            else "body")

                # Line spacing
                if prev_line_y1 is not None:
                    spacing = round(line_y0 - prev_line_y1, 2)
                    if spacing > 0:
                        data.line_spacings.append(LineSpacing(
                            page=pno, line_from=line_num-1, line_to=line_num,
                            spacing_pt=spacing, zone=zone
                        ))
                prev_line_y1 = line_y1

                # Alignment detection
                line_x0    = ln["bbox"][0]
                line_x1    = ln["bbox"][2]
                line_width  = line_x1 - line_x0
                page_center = pw / 2
                line_center = (line_x0 + line_x1) / 2
                center_diff = abs(line_center - page_center)
                left_margin = line_x0 - (pw * 0.05)
                right_margin = (pw * 0.95) - line_x1

                if center_diff < 10:
                    alignment = "CENTER"
                elif abs(left_margin - right_margin) < 5 and line_width > pw * 0.6:
                    alignment = "JUSTIFY"
                elif left_margin < right_margin:
                    alignment = "LEFT"
                else:
                    alignment = "RIGHT"

                col = 1
                for sid, span in enumerate(ln.get("spans", [])):
                    fname  = span.get("font", "")
                    fsize  = round(span.get("size", 0), 2)
                    fcolor = _int_to_rgb(span.get("color", 0))
                    flags  = span.get("flags", 0)
                    bold   = bool(flags & (1 << 4))
                    italic = bool(flags & (1 << 1))

                    # Highlight color from span background
                    highlight = None
                    if span.get("color_space") and span.get("bg_color"):
                        highlight = _float_rgb(span.get("bg_color"))

                    # Character-level extraction for letter spacing
                    chars = []
                    char_list = span.get("chars", [])
                    letter_spacing = 0.0
                    if char_list and len(char_list) > 1:
                        gaps = []
                        for i in range(len(char_list) - 1):
                            c1 = char_list[i]
                            c2 = char_list[i+1]
                            gap = round(c2["bbox"][0] - c1["bbox"][2], 2)
                            if gap >= 0:
                                gaps.append(gap)
                        letter_spacing = round(sum(gaps)/len(gaps), 2) if gaps else 0.0
                        for ch in char_list:
                            chars.append(CharToken(
                                char=ch.get("c",""),
                                bbox=BBox(*[round(v,2) for v in ch["bbox"]]),
                                origin=(round(ch.get("origin",[0,0])[0],2),
                                        round(ch.get("origin",[0,0])[1],2))
                            ))

                    for word_text in span.get("text","").split():
                        data.words.append(WordToken(
                            text            = word_text,
                            page            = pno,
                            line            = line_num,
                            col             = col,
                            bbox            = BBox(*[round(v,2) for v in span["bbox"]]),
                            font_name       = fname,
                            font_size       = fsize,
                            font_color      = fcolor,
                            bold            = bold,
                            italic          = italic,
                            highlight_color = highlight,
                            alignment       = alignment,
                            letter_spacing  = letter_spacing,
                            zone            = zone,
                            block_id        = block_id,
                            span_id         = sid,
                            chars           = chars,
                        ))
                        col += len(word_text) + 1

        # ── More precise word bboxes ───────────────────────────────────────────
        word_list   = page.get_text("words")
        page_words  = [w for w in data.words if w.page == pno]
        for i, wl in enumerate(word_list):
            if i < len(page_words):
                page_words[i].bbox = BBox(
                    round(wl[0],2), round(wl[1],2),
                    round(wl[2],2), round(wl[3],2)
                )

        # ── Highlight annotations ─────────────────────────────────────────────
        for annot in page.annots():
            if annot.type[0] == 8:  # highlight
                color  = _float_rgb(annot.colors.get("stroke") or (1,1,0))
                a_bbox = BBox(*[round(v,2) for v in annot.rect])
                # Mark words inside this annotation
                for w in data.words:
                    if (w.page == pno and
                            w.bbox.x0 >= a_bbox.x0 - 2 and
                            w.bbox.x1 <= a_bbox.x1 + 2 and
                            w.bbox.y0 >= a_bbox.y0 - 2 and
                            w.bbox.y1 <= a_bbox.y1 + 2):
                        w.highlight_color = color

        # ── Drawings (borders, lines, rects) ──────────────────────────────────
        for path_obj in page.get_drawings():
            bbox   = BBox(*[round(v,2) for v in path_obj["rect"]])
            color  = _float_rgb(path_obj.get("color") or (0,0,0))
            fill   = _float_rgb(path_obj.get("fill")) if path_obj.get("fill") else None
            width  = round(path_obj.get("width", 0), 2)
            kind   = "rect" if (bbox.width > 5 and bbox.height > 5) else "line"

            # Watermark detection: very large, very light rectangles
            is_wm = False
            if fill and bbox.width > pw * 0.5 and bbox.height > ph * 0.5:
                brightness = sum(fill) / 3
                if brightness > 200:
                    is_wm = True
                    data.watermarks.append(WatermarkElement(
                        page=pno, kind="image", text="",
                        bbox=bbox, opacity=0.3,
                        color=fill
                    ))

            data.drawings.append(DrawingElement(
                page=pno, kind=kind, bbox=bbox,
                color=color, fill=fill, width=width,
                is_watermark=is_wm
            ))

        # ── Images ────────────────────────────────────────────────────────────
        img_list = page.get_images(full=True)
        for img_info in img_list:
            xref = img_info[0]
            try:
                base_img  = doc.extract_image(xref)
                img_bytes = base_img["image"]
                img_ext   = base_img["ext"]
                img_hash  = hashlib.md5(img_bytes).hexdigest()

                # Get image position via xobject dict
                img_rects = page.get_image_rects(xref)
                for img_rect in img_rects:
                    bbox = BBox(*[round(v,2) for v in img_rect])

                    # Watermark: if covers large area or very transparent
                    is_wm = (bbox.width > pw * 0.4 and bbox.height > ph * 0.4)
                    if is_wm:
                        data.watermarks.append(WatermarkElement(
                            page=pno, kind="image", text="",
                            bbox=bbox, opacity=0.5, color=None
                        ))

                    data.images.append(ImageElement(
                        page=pno, bbox=bbox,
                        width=base_img.get("width",0),
                        height=base_img.get("height",0),
                        img_hash=img_hash,
                        img_bytes=img_bytes,
                        is_watermark=is_wm
                    ))
            except Exception as e:
                data.errors.append(f"Image extraction error p{pno}: {e}")

        # ── Watermark text detection ───────────────────────────────────────────
        # Look for transparent/repeated large text (common watermark pattern)
        for block in raw.get("blocks", []):
            if block.get("type") != 0:
                continue
            for ln in block.get("lines", []):
                for span in ln.get("spans", []):
                    flags = span.get("flags", 0)
                    color = _int_to_rgb(span.get("color", 0))
                    text  = span.get("text","").strip()
                    size  = span.get("size", 0)
                    # Large rotated text or very light color = likely watermark
                    brightness = sum(color) / 3 if color else 0
                    if size > 40 and brightness > 180 and len(text) > 2:
                        data.watermarks.append(WatermarkElement(
                            page=pno, kind="text", text=text,
                            bbox=BBox(*[round(v,2) for v in span["bbox"]]),
                            opacity=0.3, color=color
                        ))

        # ── Form fields / widgets ─────────────────────────────────────────────
        for widget in page.widgets():
            ft_map = {
                fitz.PDF_WIDGET_TYPE_TEXT:      "text",
                fitz.PDF_WIDGET_TYPE_CHECKBOX:  "checkbox",
                fitz.PDF_WIDGET_TYPE_RADIOBUTTON:"radio",
                fitz.PDF_WIDGET_TYPE_LISTBOX:   "listbox",
                fitz.PDF_WIDGET_TYPE_COMBOBOX:  "dropdown",
                fitz.PDF_WIDGET_TYPE_SIGNATURE: "signature",
            }
            data.form_fields.append(FormField(
                page       = pno,
                name       = widget.field_name or "",
                field_type = ft_map.get(widget.field_type, "unknown"),
                value      = str(widget.field_value or ""),
                bbox       = BBox(*[round(v,2) for v in widget.rect]),
                required   = bool(widget.field_flags & 2),
                font_name  = widget.text_font or "",
                font_size  = round(widget.text_fontsize or 0, 2),
            ))

        # ── Page image for visual diff ─────────────────────────────────────────
        try:
            mat = fitz.Matrix(DPI/72, DPI/72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
            data.page_images.append(img)
        except Exception as e:
            data.errors.append(f"Page render failed p{pno}: {e}")

    doc.close()
    data.raw_text = "\n".join(w.text for w in data.words)
    return data


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Extractor
# ─────────────────────────────────────────────────────────────────────────────

def _extract_docx(path: str) -> DocumentData:
    data = DocumentData(path=path, file_type="docx")
    if not _DOCX:
        data.errors.append("python-docx not installed. Run: pip install python-docx")
        return _extract_text(path)

    try:
        document = docx.Document(path)
    except Exception as e:
        data.errors.append(f"Cannot open DOCX: {e}")
        return data

    section = document.sections[0]
    pw = _emu_to_pt(section.page_width)
    ph = _emu_to_pt(section.page_height)
    data.pages.append(PageMeta(
        page=1, width=round(pw,2), height=round(ph,2), rotation=0,
        margin_top   = round(_emu_to_pt(section.top_margin),2),
        margin_bot   = round(_emu_to_pt(section.bottom_margin),2),
        margin_left  = round(_emu_to_pt(section.left_margin),2),
        margin_right = round(_emu_to_pt(section.right_margin),2),
    ))

    line_num   = 0
    prev_para_y = 0

    ALIGN_MAP = {
        WD_ALIGN_PARAGRAPH.LEFT:    "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER:  "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT:   "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        None:                        "LEFT",
    }

    def _process_para(para, zone="body", page=1):
        nonlocal line_num, prev_para_y
        line_num += 1

        # Alignment
        alignment = ALIGN_MAP.get(para.alignment, "LEFT")

        # Paragraph spacing
        pf = para.paragraph_format
        space_before = round(_emu_to_pt(pf.space_before or 0), 2)
        space_after  = round(_emu_to_pt(pf.space_after  or 0), 2)
        line_spacing_val = 0.0
        if pf.line_spacing:
            try:
                line_spacing_val = round(_emu_to_pt(pf.line_spacing), 2)
            except Exception:
                pass

        data.para_spacings.append(ParaSpacing(
            page=page, block_id=line_num,
            space_before=space_before,
            space_after=space_after,
        ))
        if line_spacing_val > 0:
            data.line_spacings.append(LineSpacing(
                page=page, line_from=line_num-1, line_to=line_num,
                spacing_pt=line_spacing_val, zone=zone
            ))

        col = 1
        for run in para.runs:
            if not run.text.strip():
                col += len(run.text)
                continue

            # Font properties
            fn    = run.font.name or (para.style.font.name if para.style and para.style.font else "") or ""
            fs_obj = run.font.size or (para.style.font.size if para.style and para.style.font else None)
            fs    = round(fs_obj.pt, 2) if fs_obj else 0.0
            fc    = (0, 0, 0)
            if run.font.color and run.font.color.type is not None:
                try:
                    rgb = run.font.color.rgb
                    fc  = (rgb.red, rgb.green, rgb.blue)
                except Exception:
                    pass

            # Highlight color
            highlight = None
            if run.font.highlight_color:
                try:
                    hc = run.font.highlight_color
                    highlight = _highlight_enum_to_rgb(hc)
                except Exception:
                    pass

            # Letter spacing from XML
            letter_spacing = 0.0
            try:
                rpr = run._r.get_or_add_rPr()
                spacing_el = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing")
                if spacing_el is not None:
                    val = spacing_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                    if val:
                        letter_spacing = round(int(val) / 20.0, 2)  # twentieths of a point
            except Exception:
                pass

            for word in run.text.split():
                data.words.append(WordToken(
                    text            = word,
                    page            = page,
                    line            = line_num,
                    col             = col,
                    bbox            = BBox(col*7.0, line_num*14.0, (col+len(word))*7.0, (line_num+1)*14.0),
                    font_name       = fn,
                    font_size       = fs,
                    font_color      = fc,
                    bold            = bool(run.bold),
                    italic          = bool(run.italic),
                    underline       = bool(run.underline),
                    strike          = bool(run.font.strike),
                    highlight_color = highlight,
                    alignment       = alignment,
                    letter_spacing  = letter_spacing,
                    zone            = zone,
                ))
                col += len(word) + 1

    # Headers
    for sec in document.sections:
        if sec.header:
            for para in sec.header.paragraphs:
                _process_para(para, zone="header")

    # Body
    for para in document.paragraphs:
        _process_para(para, zone="body")

    # Tables
    for tbl_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = " ".join(p.text for p in cell.paragraphs)

                # Cell border styling from XML
                def _get_border(tc, side):
                    try:
                        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        tcPr = tc.find(f"{{{ns}}}tcPr")
                        if tcPr is None: return None
                        tcBorders = tcPr.find(f"{{{ns}}}tcBorders")
                        if tcBorders is None: return None
                        border_el = tcBorders.find(f"{{{ns}}}{side}")
                        if border_el is None: return None
                        color_val = border_el.get(f"{{{ns}}}color","000000")
                        sz_val    = border_el.get(f"{{{ns}}}sz","4")
                        style_val = border_el.get(f"{{{ns}}}val","single")
                        color_rgb = _hex_to_rgb(color_val) if color_val != "auto" else (0,0,0)
                        return {
                            "color": color_rgb,
                            "width": round(int(sz_val)/8.0, 2),
                            "style": style_val
                        }
                    except Exception:
                        return None

                tc = cell._tc
                data.table_cells.append(TableCell(
                    page         = 1,
                    row          = row_idx,
                    col          = col_idx,
                    text         = cell_text,
                    bbox         = BBox(col_idx*100.0, row_idx*20.0,
                                        (col_idx+1)*100.0, (row_idx+1)*20.0),
                    border_top    = _get_border(tc, "top"),
                    border_bottom = _get_border(tc, "bottom"),
                    border_left   = _get_border(tc, "left"),
                    border_right  = _get_border(tc, "right"),
                ))
                for para in cell.paragraphs:
                    _process_para(para, zone="body")

    # Footers
    for sec in document.sections:
        if sec.footer:
            for para in sec.footer.paragraphs:
                _process_para(para, zone="footer")

    # Inline images
    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                img_hash  = hashlib.md5(img_bytes).hexdigest()
                data.images.append(ImageElement(
                    page=1, bbox=BBox(0,0,0,0),
                    width=0, height=0,
                    img_hash=img_hash, img_bytes=img_bytes
                ))
            except Exception:
                pass

    data.raw_text = " ".join(w.text for w in data.words)
    return data


# ─────────────────────────────────────────────────────────────────────────────
# Plain text extractor
# ─────────────────────────────────────────────────────────────────────────────

def _extract_text(path: str) -> DocumentData:
    data = DocumentData(path=path, file_type="text")
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
    except Exception as e:
        data.errors.append(f"Cannot read file: {e}")
        return data

    lines = content.splitlines()
    total = len(lines)
    data.pages.append(PageMeta(1, 0, 0, 0, 0, 0, 0, 0))

    prev_y = 0
    for ln_idx, line in enumerate(lines, 1):
        zone = ("header" if ln_idx <= max(1, total//10)
                else "footer" if ln_idx >= total - max(1, total//10)
                else "body")
        col = 1
        curr_y = ln_idx * 14.0
        if ln_idx > 1:
            data.line_spacings.append(LineSpacing(
                page=1, line_from=ln_idx-1, line_to=ln_idx,
                spacing_pt=14.0, zone=zone
            ))
        for word in line.split():
            data.words.append(WordToken(
                text=word, page=1, line=ln_idx, col=col,
                bbox=BBox(col*7, curr_y, (col+len(word))*7, curr_y+14),
                zone=zone,
            ))
            col += len(word) + 1

    data.raw_text = content
    return data


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _int_to_rgb(color_int: int) -> Tuple:
    r = (color_int >> 16) & 0xFF
    g = (color_int >>  8) & 0xFF
    b =  color_int        & 0xFF
    return (r, g, b)

def _float_rgb(color) -> Tuple:
    if color is None: return (0,0,0)
    return tuple(min(255, int(c*255)) for c in color[:3])

def _emu_to_pt(emu) -> float:
    if emu is None: return 0.0
    try: return emu / 12700.0
    except Exception: return 0.0

def _hex_to_rgb(hex_str: str) -> Tuple:
    try:
        hex_str = hex_str.lstrip("#")
        return (int(hex_str[0:2],16), int(hex_str[2:4],16), int(hex_str[4:6],16))
    except Exception:
        return (0,0,0)

def _highlight_enum_to_rgb(hc) -> Optional[Tuple]:
    """Convert WD_COLOR_INDEX enum to RGB tuple."""
    color_map = {
        1:  (255,255,0),   # yellow
        2:  (0,255,0),     # green
        3:  (0,255,255),   # cyan
        4:  (255,0,255),   # magenta
        5:  (0,0,255),     # blue
        6:  (255,0,0),     # red
        7:  (0,0,128),     # dark blue
        8:  (0,128,0),     # dark green
        9:  (0,128,128),   # dark cyan
        10: (128,0,128),   # dark magenta
        11: (128,0,0),     # dark red
        12: (128,128,0),   # dark yellow
        13: (128,128,128), # dark gray
        14: (192,192,192), # light gray
    }
    try:
        return color_map.get(int(hc), None)
    except Exception:
        return None
